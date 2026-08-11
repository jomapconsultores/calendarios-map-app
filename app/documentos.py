# ------------------------------------------------------------
# Desarrollado por Marco Antonio Posligua San Martín
# ------------------------------------------------------------
"""Lectura de los archivos que se importan al directorio: Excel, PDF y Word.

Cada formato se lee a su manera pero todos salen con la MISMA forma:

    {'tipo', 'cabeceras', 'filas', 'bloques'}

  * `filas` es la lectura tabular (Excel, y las tablas de un Word). Cuando el
    archivo trae cabeceras reconocibles se puede mapear sin gastar una llamada a
    la IA — ese es el camino barato.
  * `bloques` es texto corrido (PDF, párrafos de Word, y también las filas
    serializadas). Es lo que se le manda a la IA cuando el mapeo directo no da.

Las tres librerías se importan de forma protegida: si el servidor no tiene
`pypdf` instalado, los Excel se siguen importando y sólo el PDF avisa que le
falta la dependencia.
"""
import io
import re

try:
    import openpyxl
    EXCEL_DISPONIBLE = True
except ImportError:                                   # pragma: no cover
    EXCEL_DISPONIBLE = False

try:
    from pypdf import PdfReader
    PDF_DISPONIBLE = True
except ImportError:                                   # pragma: no cover
    PDF_DISPONIBLE = False

try:
    import docx                                       # python-docx
    WORD_DISPONIBLE = True
except ImportError:                                   # pragma: no cover
    WORD_DISPONIBLE = False


EXTENSIONES = {
    '.xlsx': 'excel', '.xlsm': 'excel', '.xls': 'excel',
    '.csv':  'csv',
    '.pdf':  'pdf',
    '.docx': 'word', '.doc': 'word',
    # Fotos: lo que de verdad manda un cliente cuando le pides sus datos —una
    # ficha fotografiada con el móvil—. Se leen con OCR (ver app/ia.py).
    '.jpg': 'imagen', '.jpeg': 'imagen', '.png': 'imagen', '.webp': 'imagen',
    '.heic': 'imagen', '.heif': 'imagen', '.bmp': 'imagen',
    '.tif': 'imagen', '.tiff': 'imagen', '.gif': 'imagen',
}

# Extensiones que se aceptan en el selector de archivos del navegador.
ACEPTA_NAVEGADOR = ','.join(sorted(EXTENSIONES))

# Cabeceras que se reconocen sin IA. La clave es lo que se lee en el archivo
# (en minúsculas y sin tildes), el valor es la columna de `contacts`.
ALIAS_CABECERAS = {
    'cedula': 'doc_number', 'cedula/ruc': 'doc_number', 'ruc': 'doc_number',
    'ci': 'doc_number', 'identificacion': 'doc_number', 'documento': 'doc_number',
    'numero de documento': 'doc_number', 'ruc/ci': 'doc_number', 'dni': 'doc_number',
    'pasaporte': 'doc_number',
    'tipo': 'doc_type', 'tipo de documento': 'doc_type', 'tipo documento': 'doc_type',
    'nombre': 'first_name', 'nombres': 'first_name', 'nombre(s)': 'first_name',
    'apellido': 'last_name', 'apellidos': 'last_name', 'apellido(s)': 'last_name',
    'nombre completo': '_nombre_completo', 'nombres y apellidos': '_nombre_completo',
    'razon social': 'business_name', 'empresa': 'business_name', 'compania': 'business_name',
    'nombre comercial': 'trade_name',
    'celular': 'mobile', 'movil': 'mobile', 'telefono celular': 'mobile', 'cel': 'mobile',
    'whatsapp': 'mobile',
    'convencional': 'landline', 'telefono': 'landline', 'telefono fijo': 'landline',
    'fijo': 'landline', 'telf': 'landline', 'tlf': 'landline',
    'correo': 'email', 'email': 'email', 'e-mail': 'email', 'correo electronico': 'email',
    'web': 'website', 'sitio web': 'website', 'pagina web': 'website',
    'direccion web': 'website', 'url': 'website',
    'direccion': 'work_address', 'direccion de trabajo': 'work_address',
    'direccion trabajo': 'work_address', 'oficina': 'work_address',
    'direccion domicilio': 'home_address', 'domicilio': 'home_address',
    'ciudad': 'city', 'canton': 'city',
    'provincia': 'province',
    'sector': 'sector', 'servicio': 'sector', 'origen': 'sector',
    'observaciones': 'notes', 'notas': 'notes', 'comentario': 'notes', 'comentarios': 'notes',
    'facebook': '_red_facebook', 'instagram': '_red_instagram',
    'linkedin': '_red_linkedin', 'twitter': '_red_x', 'x': '_red_x',
    'tiktok': '_red_tiktok', 'youtube': '_red_youtube',
    'red social': '_red_generica', 'redes sociales': '_red_generica', 'redes': '_red_generica',
}

# Cuando la cabecera no coincide exactamente con ninguna de arriba se busca por
# contenido: "Telf Celular", "N° de celular" y "CELULAR 1" son todas el celular.
# El ORDEN importa — "telefono celular" contiene «telefono» y «celular», y lo que
# manda es «celular», así que las entradas más específicas van primero.
ALIAS_POR_CONTENIDO = [
    ('celular', 'mobile'), ('movil', 'mobile'), ('whatsapp', 'mobile'),
    ('convencional', 'landline'), ('fijo', 'landline'),
    ('telefono', 'landline'), ('telf', 'landline'), ('tlf', 'landline'),
    ('ruc', 'doc_number'), ('cedula', 'doc_number'), ('identificacion', 'doc_number'),
    ('pasaporte', 'doc_number'), ('documento', 'doc_number'),
    ('razon social', 'business_name'), ('empresa', 'business_name'),
    ('nombre comercial', 'trade_name'),
    ('nombre completo', '_nombre_completo'), ('nombres y apellidos', '_nombre_completo'),
    ('apellido', 'last_name'), ('nombre', 'first_name'),
    ('correo', 'email'), ('mail', 'email'),
    ('pagina web', 'website'), ('sitio web', 'website'), ('web', 'website'),
    ('direccion de trabajo', 'work_address'), ('direccion trabajo', 'work_address'),
    ('oficina', 'work_address'),
    ('domicilio', 'home_address'),
    ('direccion', 'work_address'),
    ('provincia', 'province'), ('ciudad', 'city'), ('canton', 'city'),
    ('facebook', '_red_facebook'), ('instagram', '_red_instagram'),
    ('linkedin', '_red_linkedin'), ('tiktok', '_red_tiktok'),
    ('youtube', '_red_youtube'), ('twitter', '_red_x'),
    ('red social', '_red_generica'), ('redes', '_red_generica'),
    ('observacion', 'notes'), ('nota', 'notes'), ('comentario', 'notes'),
    ('sector', 'sector'), ('servicio', 'sector'),
]

_TILDES = str.maketrans('áéíóúüñÁÉÍÓÚÜÑ', 'aeiouunAEIOUUN')


class FormatoNoSoportado(Exception):
    """El archivo no es Excel, PDF ni Word — o falta su librería en el servidor."""


def normalizar_cabecera(texto):
    """'Teléfono Celular ' → 'telefono celular'. Sin tildes, sin espacios de más."""
    limpio = str(texto or '').translate(_TILDES).strip().lower()
    return re.sub(r'\s+', ' ', limpio)


def detectar_tipo(nombre_archivo):
    nombre = (nombre_archivo or '').lower()
    for extension, tipo in EXTENSIONES.items():
        if nombre.endswith(extension):
            return tipo
    return None


# ============================================================
#  LECTORES POR FORMATO
# ============================================================
def _leer_excel(contenido):
    if not EXCEL_DISPONIBLE:
        raise FormatoNoSoportado('Falta openpyxl en el servidor para leer archivos Excel')
    libro = openpyxl.load_workbook(io.BytesIO(contenido), data_only=True, read_only=True)
    cabeceras, filas = [], []
    for hoja in libro.worksheets:
        iterador = hoja.iter_rows(values_only=True)
        primera = next(iterador, None)
        if primera is None:
            continue
        # La primera fila con contenido se toma como cabecera de esa hoja.
        cabeceras_hoja = [str(c).strip() if c is not None else '' for c in primera]
        if not cabeceras:
            cabeceras = cabeceras_hoja
        for fila in iterador:
            if fila is None or all(c is None or str(c).strip() == '' for c in fila):
                continue
            valores = ['' if c is None else str(c).strip() for c in fila]
            filas.append(dict(zip(cabeceras_hoja, valores)))
    libro.close()
    return cabeceras, filas


def _leer_csv(contenido):
    import csv
    # Los CSV exportados desde Excel en Ecuador suelen venir en latin-1 y con ';'.
    for codificacion in ('utf-8-sig', 'utf-8', 'latin-1'):
        try:
            texto = contenido.decode(codificacion)
            break
        except UnicodeDecodeError:
            continue
    else:
        raise FormatoNoSoportado('No se pudo leer la codificación del CSV')
    muestra = texto[:4096]
    try:
        dialecto = csv.Sniffer().sniff(muestra, delimiters=',;\t|')
    except csv.Error:
        dialecto = csv.excel
    lector = csv.DictReader(io.StringIO(texto), dialect=dialecto)
    filas = [{k: (v or '').strip() for k, v in fila.items() if k}
             for fila in lector]
    return list(lector.fieldnames or []), filas


def _leer_pdf(contenido):
    if not PDF_DISPONIBLE:
        raise FormatoNoSoportado('Falta pypdf en el servidor para leer archivos PDF')
    lector = PdfReader(io.BytesIO(contenido))
    bloques = []
    for numero, pagina in enumerate(lector.pages, 1):
        try:
            texto = pagina.extract_text() or ''
        except Exception:
            texto = ''
        for linea in texto.splitlines():
            linea = linea.strip()
            if linea:
                bloques.append(linea)
        if numero >= 200:      # tope de seguridad para PDF enormes
            break
    # Sin texto seleccionable el PDF es una imagen: lo resuelve el OCR, no este
    # lector. Se avisa con una marca para que `extraer` lo derive, en vez de
    # rechazar el archivo como se hacía antes.
    return bloques


def _leer_word(contenido):
    if not WORD_DISPONIBLE:
        raise FormatoNoSoportado('Falta python-docx en el servidor para leer archivos Word')
    documento = docx.Document(io.BytesIO(contenido))
    bloques = [p.text.strip() for p in documento.paragraphs if p.text.strip()]

    # Las tablas de un Word sí son tabulares: se leen como un Excel.
    cabeceras, filas = [], []
    for tabla in documento.tables:
        if not tabla.rows:
            continue
        cabeceras_tabla = [c.text.strip() for c in tabla.rows[0].cells]
        if not cabeceras:
            cabeceras = cabeceras_tabla
        for fila in tabla.rows[1:]:
            valores = [c.text.strip() for c in fila.cells]
            if any(valores):
                filas.append(dict(zip(cabeceras_tabla, valores)))
    return cabeceras, filas, bloques


def extraer(nombre_archivo, contenido, ocr=None):
    """Lee el archivo y devuelve {'tipo', 'cabeceras', 'filas', 'bloques', 'ocr'}.

    `ocr` es una función `(contenido, nombre) -> [líneas]` que se usa para lo que
    no tiene texto extraíble: una foto, o un PDF escaneado. Se recibe como
    parámetro en lugar de importar app/ia.py aquí para que este módulo siga
    siendo sólo un lector de archivos, sin saber qué motor de IA hay detrás."""
    tipo = detectar_tipo(nombre_archivo)
    if not tipo:
        raise FormatoNoSoportado(
            'Formato no soportado. Se aceptan Excel (.xlsx), CSV, PDF, Word (.docx) '
            'y fotos (JPG, PNG, HEIC…).')

    cabeceras, filas, bloques = [], [], []
    uso_ocr = False

    if tipo == 'excel':
        cabeceras, filas = _leer_excel(contenido)
    elif tipo == 'csv':
        cabeceras, filas = _leer_csv(contenido)
    elif tipo == 'word':
        cabeceras, filas, bloques = _leer_word(contenido)
    elif tipo == 'imagen':
        # Una foto no tiene texto que extraer: o hay OCR, o no hay nada.
        if not ocr:
            raise FormatoNoSoportado(
                'Para importar una foto hace falta el OCR, que necesita '
                'MISTRAL_API_KEY configurada en el servidor.')
        bloques = ocr(contenido, nombre_archivo)
        uso_ocr = True
        if not bloques:
            raise FormatoNoSoportado('No se reconoció ningún texto en la imagen.')
    elif tipo == 'pdf':
        bloques = _leer_pdf(contenido)
        if not bloques:
            # PDF escaneado: es una imagen dentro de un PDF. Antes se rechazaba;
            # ahora se manda al OCR, que es justo el caso para el que sirve.
            if not ocr:
                raise FormatoNoSoportado(
                    'El PDF no tiene texto seleccionable (está escaneado). Para '
                    'leerlo hace falta el OCR, que necesita MISTRAL_API_KEY '
                    'configurada en el servidor.')
            bloques = ocr(contenido, nombre_archivo)
            uso_ocr = True
            if not bloques:
                raise FormatoNoSoportado('No se reconoció ningún texto en el PDF escaneado.')

    # Toda fila tabular se serializa también como bloque de texto: es lo que se
    # le entrega a la IA cuando las cabeceras no alcanzan para mapear.
    for fila in filas:
        partes = [f'{k}: {v}' for k, v in fila.items() if str(v or '').strip()]
        if partes:
            bloques.append(' | '.join(partes))

    return {'tipo': tipo, 'cabeceras': cabeceras, 'filas': filas,
            'bloques': bloques, 'ocr': uso_ocr}


# ============================================================
#  MAPEO DIRECTO POR CABECERAS (sin IA)
# ============================================================
def destino_de_cabecera(cabecera):
    """A qué columna del directorio corresponde una cabecera del archivo.

    Primero coincidencia exacta; si no, por contenido. Sin la segunda pasada,
    cabeceras tan comunes como «Telf Celular» o «N° de cédula» se descartaban en
    silencio y esos datos se perdían en la importación."""
    limpia = normalizar_cabecera(cabecera)
    if not limpia:
        return None
    if limpia in ALIAS_CABECERAS:
        return ALIAS_CABECERAS[limpia]
    for fragmento, destino in ALIAS_POR_CONTENIDO:
        if fragmento in limpia:
            return destino
    return None


def cabeceras_reconocidas(cabeceras):
    """Cuántas de las cabeceras del archivo se entienden sin IA. Sirve para
    decidir si vale la pena gastar la llamada al modelo."""
    reconocidas = {}
    for cabecera in (cabeceras or []):
        destino = destino_de_cabecera(cabecera)
        if destino:
            reconocidas[normalizar_cabecera(cabecera)] = destino
    return reconocidas


def _partir_nombre(completo):
    """'PEREZ LOPEZ, JUAN CARLOS' o 'PEREZ LOPEZ JUAN CARLOS' → (nombres, apellidos).

    Con coma es inequívoco. Sin coma se aplica la convención ecuatoriana: los dos
    primeros son apellidos."""
    texto = re.sub(r'\s+', ' ', str(completo or '').strip())
    if not texto:
        return '', ''
    if ',' in texto:
        apellidos, _, nombres = texto.partition(',')
        return nombres.strip(), apellidos.strip()
    partes = texto.split()
    if len(partes) >= 4:
        return ' '.join(partes[2:]), ' '.join(partes[:2])
    if len(partes) == 3:
        return partes[2], ' '.join(partes[:2])
    if len(partes) == 2:
        return partes[1], partes[0]
    return texto, ''


def _es_ruc_de_empresa(documento):
    """True si el número es un RUC de sociedad privada (tercer dígito 9) o de
    entidad pública (6). Sirve para no partir «COMERCIAL X S.A.» en nombres y
    apellidos: eso es una razón social, no una persona."""
    digitos = re.sub(r'\D', '', str(documento or ''))
    return len(digitos) == 13 and digitos[2] in ('6', '9')


def mapear_por_cabeceras(filas, cabeceras):
    """Convierte filas tabulares en registros del directorio usando sólo las
    cabeceras conocidas. Devuelve [] si el archivo no trae ni el documento ni un
    nombre reconocibles: en ese caso hay que ir a la IA."""
    mapa = cabeceras_reconocidas(cabeceras)
    destinos = set(mapa.values())
    if 'doc_number' not in destinos and not {'first_name', 'last_name',
                                             '_nombre_completo', 'business_name'} & destinos:
        return []

    registros = []
    for fila in filas:
        registro = {'socials': []}
        nombre_completo = ''
        for clave, valor in fila.items():
            destino = mapa.get(normalizar_cabecera(clave))
            texto = str(valor or '').strip()
            if not destino or not texto:
                continue
            if destino == '_nombre_completo':
                # Se guarda entero y se decide al final, cuando ya se sabe si el
                # documento de esta fila es de una persona o de una empresa.
                nombre_completo = texto
            elif destino.startswith('_red_'):
                red = destino.replace('_red_', '').replace('generica', 'Web').title()
                registro['socials'].append({'red': red, 'url': texto})
            else:
                registro[destino] = texto

        # Una celda de documento sin un solo dígito no es un documento: casi
        # siempre es un nombre que se coló en esa columna. Se rescata como nombre
        # en lugar de guardarse como un documento imposible.
        documento = registro.get('doc_number', '')
        if documento and not re.search(r'\d', documento):
            if not nombre_completo:
                nombre_completo = documento
            registro.pop('doc_number')
            documento = ''

        if nombre_completo:
            if _es_ruc_de_empresa(documento):
                registro.setdefault('business_name', nombre_completo)
            else:
                nombres, apellidos = _partir_nombre(nombre_completo)
                registro.setdefault('first_name', nombres)
                registro.setdefault('last_name', apellidos)

        if any(registro.get(c) for c in ('doc_number', 'first_name', 'last_name', 'business_name')):
            registros.append(registro)
    return registros
